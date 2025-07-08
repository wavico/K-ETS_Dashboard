import os
import pandas as pd
from pathlib import Path
import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO

def get_data_context() -> str:
    """
    'data' 폴더에 있는 주요 CSV 파일들을 로드하여,
    Pandas를 이용해 데이터에 대한 요약 정보를 생성하고 텍스트로 반환

    이 함수는 보고서 생성 에이전트에게 사실 기반 정보를 제공하여
    더욱 신뢰도 높은 컨텐츠를 생성

    Returns:
        str: 데이터의 주요 통계 및 구조를 요약한 텍스트.
             데이터 로드나 분석에 실패할 경우, 에러 메시지를 포함한 문자열을 반환
    """
    print("--- 데이터 컨텍스트 생성 시작 ---")
    try:
        # 현재 파일 위치를 기준으로 프로젝트 루트를 찾고, 'data' 폴더 경로를 설정
        data_folder = Path(__file__).parent / "data"

        # 분석할 주요 CSV 파일 목록
        csv_files = [
            "환경부 온실가스종합정보센터_국가 온실가스 인벤토리 배출량_20250103.csv",
            "배출권_거래데이터.csv",
            "한국에너지공단_산업부문 에너지사용 및 온실가스배출량 통계_20231231.csv"
        ]

        dataframes = []
        for filename in csv_files:
            filepath = data_folder / filename
            if filepath.exists():
                # 다양한 인코딩으로 파일 로드 시도
                for encoding in ['cp949', 'euc-kr', 'utf-8', 'utf-8-sig']:
                    try:
                        df = pd.read_csv(filepath, encoding=encoding, low_memory=False)
                        dataframes.append(df)
                        print(f"✅ '{filename}' 로드 성공 (인코딩: {encoding})")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    print(f"⚠️ '{filename}' 파일 로드 실패 (모든 인코딩 실패)")
            else:
                print(f"⚠️ '{filename}' 파일을 찾을 수 없음")
        
        if not dataframes:
            return "데이터 파일을 찾거나 로드할 수 없어 데이터 컨텍스트를 생성할 수 없습니다."

        # 모든 데이터프레임을 하나로 통합
        # sort=False 옵션은 불필요한 정렬을 방지합니다.
        combined_df = pd.concat(dataframes, ignore_index=True, sort=False)

        # 데이터 요약 정보 생성
        # StringIO를 사용해 df.info()의 출력을 문자열로 캡처
        buffer = io.StringIO()
        combined_df.info(buf=buffer)
        info_str = buffer.getvalue()

        # 수치형 데이터에 대한 기술 통계 요약
        # include='number'는 숫자 타입의 컬럼만 선택하도록 합니다.
        desc_str = combined_df.describe(include='number').to_string()

        # 최종 컨텍스트 문자열 조합
        context = f"""
### 데이터 개요 (Data Overview)
{info_str}

### 주요 수치 데이터 통계 (Key Numerical Statistics)
{desc_str}

### 데이터 샘플 (상위 5개 행)
{combined_df.head().to_string()}
"""
        print("--- ✅ 데이터 컨텍스트 생성 완료 ---")
        return context.strip()

    except Exception as e:
        error_message = f"데이터 컨텍스트 생성 중 오류가 발생했습니다: {e}"
        print(f"--- ❌ {error_message} ---")
        return error_message

if __name__ == '__main__':
    """
    이 스크립트를 직접 실행하여 get_data_context 함수의 출력을 테스트합니다.
    """
    data_summary = get_data_context()
    print("\n--- 생성된 데이터 컨텍스트 ---")
    print(data_summary)

# --- 보고서 파일 생성 유틸리티 ---

def create_docx(title: str, content: str) -> BytesIO:
    """
    보고서 제목과 내용을 받아 DOCX 파일 객체를 생성합니다.

    Args:
        title (str): 보고서의 제목.
        content (str): 보고서의 전체 본문 내용.

    Returns:
        BytesIO: 메모리 상에 생성된 DOCX 파일 객체.
    """
    document = Document()
    document.add_heading(title, level=1)
    
    # 내용을 문단별로 나누어 추가
    for paragraph in content.split('\n'):
        if paragraph.strip():
            document.add_paragraph(paragraph)
            
    # 파일 객체를 메모리에 저장
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_pdf(title: str, content: str) -> BytesIO:
    """
    보고서 제목과 내용을 받아 PDF 파일 객체를 생성합니다.
    한글을 지원하기 위해 'NanumGothic' 폰트를 사용합니다.

    Args:
        title (str): 보고서의 제목.
        content (str): 보고서의 전체 본문 내용.

    Returns:
        BytesIO: 메모리 상에 생성된 PDF 파일 객체.
    """
    # 폰트 등록 (파일 경로를 확인해야 할 수 있음, 임시로 상대 경로 사용)
    # 실제 배포 환경에서는 폰트 파일의 경로를 정확히 지정해야 합니다.
    try:
        import platform
        system = platform.system()
        
        if system == "Windows":
            font_paths = ["c:/Windows/Fonts/malgun.ttf", "c:/Windows/Fonts/gulim.ttc"]
        elif system == "Darwin":  # macOS
            font_paths = ["/System/Library/Fonts/AppleGothic.ttf"]
        else:  # Linux
            font_paths = ["/usr/share/fonts/truetype/nanum/NanumGothic.ttf"]
        
        font_name = 'Helvetica'  # fallback
        for font_path in font_paths:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('KoreanFont', font_path))
                font_name = 'KoreanFont'
                break
    except Exception:
        # 폰트 로드 실패 시 기본 폰트로 대체 (한글 깨질 수 있음)
        print("경고: '맑은 고딕' 폰트를 찾을 수 없습니다. PDF에서 한글이 깨질 수 있습니다.")
        font_name = 'Helvetica'

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # 제목
    p.setFont(font_name, 16)
    p.drawString(50, height - 50, title)

    # 내용
    p.setFont(font_name, 10)
    text = p.beginText(50, height - 100)
    text.setLeading(15)

    for line in content.split('\n'):
        text.textLine(line)

    p.drawText(text)
    p.showPage()
    p.save()
    
    buffer.seek(0)
    return buffer 

