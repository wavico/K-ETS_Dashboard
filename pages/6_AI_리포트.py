import streamlit as st
import sys
import os
import pdfplumber
from openai import OpenAI
from datetime import datetime
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pinecone import Pinecone, ServerlessSpec
from langchain.embeddings import OpenAIEmbeddings
import uuid
from langchain.chat_models import ChatOpenAI
from io import BytesIO
from docx import Document
from docx.shared import Pt

# 경로 추가
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Streamlit 페이지 설정
st.set_page_config(
    page_title="AI_리포트 생성기",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 커스텀 CSS
st.markdown("""
<style>
    .main-header { font-size: 28px; font-weight: bold; margin-bottom: 20px; }
</style>
""", unsafe_allow_html=True)

# 타이틀
st.markdown('<h1 class="main-header">📄 AI 기반 보고서 생성기</h1>', unsafe_allow_html=True)

# API 키 로딩
load_dotenv()

# vLLM 설정 확인
use_vllm = os.getenv("USE_VLLM", "true").lower() == "true"
vllm_base_url = os.getenv("VLLM_BASE_URL", "http://localhost:8000/v1")

if use_vllm:
    # vLLM 사용 시
    client = OpenAI(
        base_url=vllm_base_url,
        api_key="EMPTY"  # vLLM은 API 키 불필요
    )
    os.environ["OPENAI_API_KEY"] = "EMPTY"
else:
    # 기존 OpenAI API 사용
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

pinecone_api_key = os.getenv("PINECONE_API_KEY")
index_name = "carbone-index"
pc = Pinecone(api_key=pinecone_api_key)

# PDF 텍스트 추출
@st.cache_data
def extract_text_from_pdf(uploaded_file):
    with pdfplumber.open(uploaded_file) as pdf:
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

# 목차 추출
def extract_table_of_contents(text):
    prompt = f"""
다음 문서에서 **목차(차례)**에 해당하는 부분만 정확히 추출해 주세요.
{text[:4000]}
"""
    model_name = os.getenv("VLLM_MODEL_NAME", "gpt-4-turbo") if use_vllm else "gpt-4-turbo"
    response = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": "당신은 문서 구조에서 목차만 정확히 추출하는 AI입니다."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

# 문서 형식 요약
def summarize_template_structure(text):
    prompt = f"""
다음 문서의 형식(보고서 구조, 제목 스타일, 구성 흐름 등)을 간단히 요약해 주세요.
{text[:4000]}
"""
    model_name = os.getenv("VLLM_MODEL_NAME", "gpt-4-turbo") if use_vllm else "gpt-4-turbo"
    response = client.chat.completions.create(
        model=model_name,
        messages=[
            {"role": "system", "content": "당신은 문서 형식을 분석하고 요약하는 AI입니다."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()

# Pinecone 기반 벡터 저장소 생성
@st.cache_resource
def create_vector_store(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    docs = text_splitter.create_documents([text])
    embeddings_model = OpenAIEmbeddings()
    vectors = embeddings_model.embed_documents([doc.page_content for doc in docs])

    if index_name not in pc.list_indexes().names():
        pc.create_index(
            name=index_name,
            dimension=1536,
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1")
        )

    index = pc.Index(index_name)
    to_upsert = []
    for i, vector in enumerate(vectors):
        to_upsert.append({
            "id": str(uuid.uuid4()),
            "values": vector,
            "metadata": {"text": docs[i].page_content}
        })
    index.upsert(vectors=to_upsert)
    return index

# 유사 문서 검색
def retrieve_similar_docs(topic, index, embeddings_model, top_k=5):
    embedded_query = embeddings_model.embed_query(topic)
    result = index.query(vector=embedded_query, top_k=top_k, include_metadata=True)
    return [match["metadata"]["text"] for match in result["matches"]]

# 보고서 생성
def generate_report_with_rag(topic, index, custom_outline=None):
    model_name = os.getenv("VLLM_MODEL_NAME", "gpt-4-turbo") if use_vllm else "gpt-4-turbo"

    if use_vllm:
        llm = ChatOpenAI(
            base_url=vllm_base_url,
            api_key="EMPTY",
            temperature=0.7,
            model_name=model_name
        )
    else:
        llm = ChatOpenAI(temperature=0.7, model_name="gpt-4-turbo")

    embeddings_model = OpenAIEmbeddings()
    docs = retrieve_similar_docs(topic, index, embeddings_model)
    context = "\n\n".join(docs[:5])

    outline_part = f"\n다음 목차 구조를 따르세요:\n{custom_outline}" if custom_outline else ""

    prompt = f"""
'{topic}'에 대한 보고서를 작성해 주세요.
- 참고 문서 기반 형식을 따르되, 복사하지 말고 새로운 내용으로 작성하세요.
📚 참고 문서:
{context}{outline_part}
"""
    response = llm.invoke(prompt)
    return response.content.strip()

# Word 문서 생성
def generate_docx_report(text, topic):
    doc = Document()
    doc.add_heading(topic, level=1)
    for paragraph in text.split("\n"):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            p.style.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.5
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 사용자 입력 순서
st.subheader("📝 Step 1: 주제 입력")
topic = st.text_input("보고서 주제를 입력하세요", placeholder="예: 탄소중립 추진 전략")

st.subheader("📎 Step 2: 템플릿 PDF 업로드")
uploaded_file = st.file_uploader("템플릿 문서를 업로드하세요", type=["pdf"])

custom_outline = ""
if uploaded_file:
    with st.spinner("문서 분석 중..."):
        extracted_text = extract_text_from_pdf(uploaded_file)
        toc = extract_table_of_contents(extracted_text)
        structure_summary = summarize_template_structure(extracted_text)
        vector_store = create_vector_store(extracted_text)

    st.subheader("📚 문서 형식 요약")
    st.markdown(structure_summary)

    st.subheader("📑 추출된 템플릿 목차")
    st.code(toc, language="markdown")

    st.subheader("✏️ Step 3: 사용할 목차를 수정 또는 그대로 사용")
    custom_outline = st.text_area("사용할 목차를 입력 또는 편집하세요", value=toc, height=200)

# Step 4: 생성 실행
if topic and uploaded_file and custom_outline:
    with st.spinner("📄 새로운 보고서를 작성하는 중입니다..."):
        generated_report = generate_report_with_rag(topic, vector_store, custom_outline)
        docx_file = generate_docx_report(generated_report, topic)

    st.success("✅ 보고서 생성 완료!")
    st.download_button(
        label="📥 Word 형식으로 보고서 다운로드 (.docx)",
        data=docx_file,
        file_name=f"{topic}_보고서.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.text_area("📄 생성된 보고서 미리보기", generated_report, height=500)
